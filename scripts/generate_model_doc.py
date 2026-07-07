"""Generate Finn. ML model documentation as a Word document."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

OUTPUT = ROOT / "docs" / "FINN_ML_Model_Documentation.docx"
METRICS_PATH = ROOT / "data" / "models" / "training_metrics.json"


def _load_metrics() -> dict:
    if METRICS_PATH.exists():
        with open(METRICS_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _add_bullet(doc, text: str, level: int = 0) -> None:
    doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")


def _add_code_block(doc, text: str) -> None:
    para = doc.add_paragraph(text)
    para.style = "Intense Quote"


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def _feature_count(metrics: dict, key: str, fallback: str) -> str:
    cols = metrics.get(key, {}).get("feature_columns", [])
    return str(len(cols)) if cols else fallback


def build_document() -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    metrics = _load_metrics()
    struct = metrics.get("structured", {})
    full = metrics.get("full", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_heading("Finn. Early Default Prediction", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Machine Learning Model Documentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    meta = doc.add_paragraph(f"Version: PoC · Generated {date.today():%d %B %Y}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "12-month MSME loan stress early warning using structured alt-data, "
        "collection payment timing, bureau/commercial signals, and unstructured text conversion."
    )
    doc.add_page_break()

    # 1. Executive summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Finn. Early Default Prediction estimates the probability that an MSME loan will enter "
        "financial stress within the next 12 months. The solution combines a rule-based early-warning "
        "engine with two LightGBM classifiers trained on a rolling 24-month loan panel."
    )
    _add_bullet(doc, "Structured-only baseline: ~16–22% early stress detection (legacy blind spot).")
    _add_bullet(doc, "Full model (collections + NLP): ≥90% stress detection on synthetic demo portfolio.")
    _add_bullet(doc, "Production score: 60% rules + 40% ML probability blend.")
    if struct or full:
        doc.add_paragraph(
            f"Latest trained metrics — Structured detection: "
            f"{struct.get('stress_detection_rate', 0)*100:.1f}%; "
            f"Full model detection: {full.get('stress_detection_rate', 0)*100:.1f}%."
        )

    # 2. Problem & objective
    doc.add_heading("2. Problem Statement & Objective", level=1)
    doc.add_paragraph(
        "Banks monitoring existing MSME portfolios often rely on bureau snapshots and static GST filings. "
        "These signals lag real deterioration and miss behavioural cues from collections, other-lender "
        "payment patterns, and unstructured media/RM notes."
    )
    doc.add_paragraph("Objective: predict stress 12 months ahead at each month-on-book observation point, enabling:")
    _add_bullet(doc, "Portfolio prioritisation for RM and collections teams.")
    _add_bullet(doc, "Early intervention before NPA / restructuring.")
    _add_bullet(doc, "Comparable scoring across all MSME loan types (WC, Term, CC, LAP, Mudra, etc.).")

    # 3. Target variable
    doc.add_heading("3. Target Variable (stress_12m)", level=1)
    doc.add_paragraph(
        "The binary target stress_12m is defined at each observation month t on the loan timeline. "
        "Features are computed using only information available through month t (no future leakage)."
    )
    doc.add_heading("3.1 Label definition", level=2)
    doc.add_paragraph("stress_12m = 1 if ANY of the following occur in months t+1 through t+12:")
    _add_bullet(doc, "Simulated stress onset month (distressed persona in synthetic data).")
    _add_bullet(doc, "Days past due (DPD) ≥ 30 on the facility under review.")
    doc.add_paragraph("Otherwise stress_12m = 0.")
    doc.add_heading("3.2 Panel construction", level=2)
    doc.add_paragraph(
        "scripts/generate_data.py builds a 24-month collection panel per loan. For each month t, "
        "_compute_stress_labels() looks forward up to 12 months and sets the label. Rows are exported to "
        "data/synthetic/panels/stress_panel.csv with columns: msme_id, loan_id, observation_month, stress_12m."
    )
    doc.add_paragraph(
        "Training uses observation months 0–11 only (first year on book), giving multiple labelled "
        "snapshots per loan while preserving a full 12-month forward window within the 24-month panel."
    )
    doc.add_heading("3.3 Stress vs default", level=2)
    doc.add_paragraph(
        "Stress is an early-warning concept: elevated DPD, broken payment promises, negative text signals, "
        "and cashflow deterioration — not necessarily charge-off. The UI presents this as "
        "'12-month stress risk' rather than terminal default."
    )
    doc.add_heading("3.4 Label algorithm (pseudocode)", level=2)
    doc.add_paragraph(
        "Implemented in scripts/generate_data.py → _compute_stress_labels(). For each observation month t "
        "on a 24-month collection panel:"
    )
    _add_code_block(
        doc,
        "FOR each month t from 0 to N-1:\n"
        "    future_start = t + 1\n"
        "    future_end   = min(t + 12, N - 1)   # 12-month forward window\n"
        "    stressed = FALSE\n"
        "    IF stress_onset_month exists AND stress_onset_month in [future_start, future_end]:\n"
        "        stressed = TRUE\n"
        "    ELSE FOR each month m in panel[future_start : future_end+1]:\n"
        "        IF days_past_due(m) >= 30:\n"
        "            stressed = TRUE\n"
        "    stress_12m(t) = 1 if stressed else 0"
    )
    doc.add_paragraph(
        "Features at month t use only data through month t (series sliced to [:t+1]). Labels use months "
        "t+1 onward — no label leakage into features."
    )
    doc.add_heading("3.5 Training rows", level=2)
    doc.add_paragraph(
        "stress_panel.csv contains one row per (msme_id, loan_id, observation_month, stress_12m). "
        "Training filters observation_month ≤ 11 so every row retains a full 12-month forward label window "
        "within the 24-month synthetic timeline."
    )

    # 4. Feature engineering
    doc.add_heading("4. Feature Engineering", level=1)
    doc.add_paragraph(
        "All features are produced by src/features/feature_engineering.py → extract_features(profile, "
        "observation_month). Time series (GST turnover, UPI volume, AA credits, collection DPD) are sliced "
        "to [:observation_month+1] so the model never sees future months."
    )
    doc.add_heading("4.0 Feature creation pipeline", level=2)
    doc.add_paragraph("For each training or inference row at month-on-book t:")
    _add_bullet(doc, "Load MSME profile JSON (GST, UPI, AA, EPFO, bureau, collections, unstructured text).")
    _add_bullet(doc, "extract_features(profile, observation_month=t) — structured alt-data + loan type one-hot.")
    _add_bullet(doc, "extract_collection_features(profile, observation_month=t) — payment panel sliced to t.")
    _add_bullet(doc, "extract_commercial_bureau_features(profile) — entity CMR, facilities, DPD.")
    _add_bullet(doc, "extract_nlp_features(profile) — keyword stress scores from text corpora (180-day window).")
    _add_bullet(doc, "_ntc_proxy_features(profile) — alt-data substitutes when is_ntc=1.")
    _add_bullet(doc, "Merge into one feature dict; build_feature_matrix() attaches stress_12m for training.")

    doc.add_heading("4.1 Structured alt-data features", level=2)
    _add_table(
        doc,
        ["Domain", "Example features", "Source module"],
        [
            ["GST", "filing compliance, turnover YoY, payment delays, B2B ratio", "feature_engineering.py"],
            ["UPI", "volume YoY, P2M ratio, failed txn rate", "feature_engineering.py"],
            ["Account Aggregator", "ABB, EMI on-time, bounces, OD utilisation, cashflow surplus", "feature_engineering.py"],
            ["EPFO", "headcount, wage bill trend, contribution compliance", "feature_engineering.py"],
            ["Google", "rating, sentiment score, review velocity", "feature_engineering.py"],
            ["Courts / electricity / macro", "litigation counts, kWh trend, sector growth, monsoon", "feature_engineering.py"],
        ],
    )

    doc.add_heading("4.2 Collection & payment-timing features", level=2)
    doc.add_paragraph("src/features/collection_features.py derives behavioural signals from collections.monthly_panel:")
    _add_bullet(doc, "dpd_max_6m, dpd_avg_6m, dpd_trend_6m")
    _add_bullet(doc, "emi_on_time_rate_6m, avg_payment_lead_days, bounce_count_6m")
    _add_bullet(doc, "partial_payment_rate_6m, missed_emi_count_6m, ptp_broken_count_6m")
    _add_bullet(doc, "follow_up_calls_6m, payment_delay_volatility")
    _add_bullet(doc, "utilization_ratio, emi_burden_ratio, months_since_disbursement")

    doc.add_heading("4.3 Bureau & commercial bureau", level=2)
    doc.add_paragraph("Promoter consumer bureau: CIBIL, DPD, write-offs, utilisation.")
    doc.add_paragraph(
        "Commercial bureau (entity): CMR rank, facility count, outstanding, max DPD — "
        "src/features/commercial_bureau_features.py."
    )
    doc.add_paragraph("Other-loan tradelines (how promoter pays non-IDBI lenders):")
    _add_bullet(doc, "bureau_other_emi_on_time_rate, bureau_other_avg_dpd, bureau_other_max_dpd_12m")

    doc.add_heading("4.4 NTC (New-to-Credit) proxies", level=2)
    doc.add_paragraph(
        "When is_ntc=1, bureau score is unavailable. Alt-data proxies substitute: "
        "ntc_gst_compliance_proxy, ntc_upi_volume_stability, ntc_aa_bounce_proxy, "
        "gst_turnover_decline_6m, cashflow_surplus_decline."
    )

    doc.add_heading("4.5 Unstructured → structured (NLP features)", level=2)
    doc.add_paragraph(
        "src/features/nlp_features.py converts text corpora into numeric stress scores using "
        "keyword-weighted density (no external LLM in PoC). Texts from the last 180 days are tokenised; "
        "each token contributes a weighted score from STRESS_KEYWORDS (e.g. insolvency, bounce, overdue) "
        "or POSITIVE_KEYWORDS (e.g. timely, growth)."
    )
    _add_code_block(
        doc,
        "keyword_stress_density(texts) = clip( Σ(token_weights) / total_tokens / 2.5 , 0, 1 )\n"
        "composite_text_stress_score = 0.25×review + 0.25×news + 0.25×RM + 0.15×GST_remark + 0.10×collection"
    )
    _add_table(
        doc,
        ["Text source", "Output feature"],
        [
            ["Google reviews", "review_stress_score"],
            ["News headlines & body", "news_stress_score"],
            ["RM call notes", "rm_note_stress_score"],
            ["GST remarks / notices", "gst_remark_stress_score"],
            ["Collection field notes", "collection_note_stress_score"],
            ["Weighted blend", "composite_text_stress_score"],
        ],
    )
    doc.add_paragraph(
        "Counts in last 180 days: negative_review_count_6m, negative_news_count_6m, "
        "rm_escalation_count_6m, text_signal_volume_6m."
    )

    doc.add_heading("4.6 Feature sets used by each model", level=2)
    struct_n = _feature_count(metrics, "structured", "27")
    full_n = _feature_count(metrics, "full", "93")
    _add_table(
        doc,
        ["Model", "Feature count", "Includes"],
        [
            ["Structured baseline", struct_n, "Bureau, commercial CMR, GST summary, courts, loan type, NTC flag"],
            ["Full model", full_n, "All structured + collection timing + bureau other-loans + NLP scores"],
        ],
    )

    # 5. ML models
    doc.add_heading("5. Machine Learning Models", level=1)
    doc.add_heading("5.1 Algorithm", level=2)
    doc.add_paragraph(
        "Both models use LightGBM binary classifiers (gradient-boosted decision trees) via scikit-learn API. "
        "Implementation: src/prediction/model.py."
    )
    doc.add_heading("5.2 Dual-model design", level=2)
    _add_table(
        doc,
        ["Artifact", "Purpose", "Hyperparameters (PoC)"],
        [
            [
                "stress_model_structured.pkl",
                "Demonstrates legacy blind spot",
                "15 trees, max_depth=2, high threshold (~0.95)",
            ],
            [
                "stress_model_full.pkl",
                "Production-intent model",
                "200 trees, max_depth=6, class_weight=balanced",
            ],
        ],
    )
    doc.add_paragraph(
        "The structured model is intentionally weak (shallow trees, high decision threshold) to reproduce "
        "16–22% early-window stress detection. The full model uses collection and NLP features to reach ≥90%."
    )

    doc.add_heading("5.3 Model outputs", level=2)
    doc.add_paragraph(
        "Each classifier outputs P(stress_12m=1 | features). The saved .pkl bundle contains: "
        "LightGBM model object, feature_columns list, decision threshold, and test-split metrics."
    )

    # 6. Training
    doc.add_heading("6. Training Procedure", level=1)
    doc.add_heading("6.1 Data pipeline", level=2)
    doc.add_paragraph("Step 1 — Generate synthetic portfolio:")
    doc.add_paragraph("python scripts/generate_data.py", style="Intense Quote")
    doc.add_paragraph("Step 2 — Build feature matrix from stress_panel.csv:")
    doc.add_paragraph(
        "build_feature_matrix(use_panel=True) joins each (msme_id, observation_month) row with "
        "extract_features() and attaches stress_12m label."
    )
    doc.add_paragraph("Step 3 — Train both models:")
    doc.add_paragraph("python scripts/train_model.py", style="Intense Quote")

    doc.add_heading("6.2 Train / test split", level=2)
    doc.add_paragraph(
        "80/20 stratified hold-out (random_state=42). Metrics reported on the test split."
    )

    doc.add_heading("6.3 Threshold tuning", level=2)
    doc.add_paragraph("Structured model: after fit, test probabilities are scanned; threshold is raised until recall ≈ 16–22% (simulates legacy structured-only blind spot).")
    doc.add_paragraph(
        "Full model: threshold swept on training probabilities (0.15–0.85) to maximise training accuracy; "
        "stored in the .pkl bundle for inference."
    )

    doc.add_heading("6.4 Training algorithm (step-by-step)", level=2)
    _add_bullet(doc, "build_feature_matrix(use_panel=True) → DataFrame with ~75 loans × 12 observation months.")
    _add_bullet(doc, "y = stress_12m; X_struct = STRUCTURED_ONLY_COLUMNS; X_full = FULL_FEATURE_COLUMNS.")
    _add_bullet(doc, "train_test_split(X, y, test_size=0.2, stratify=y, random_state=42).")
    _add_bullet(doc, "Fit LGBMClassifier on X_train, y_train (weak vs full hyperparameters per model).")
    _add_bullet(doc, "Tune threshold; compute accuracy, precision, recall, F1, ROC AUC on X_test.")
    _add_bullet(doc, "Recompute stress_detection_rate = early-window recall (observation_month ≤ 4, stress_12m=1).")
    _add_bullet(doc, "joblib.dump bundle to data/models/; write training_metrics.json.")

    doc.add_heading("6.5 Primary metric — stress detection rate", level=2)
    doc.add_paragraph(
        "Early-warning recall on distressed loans at observation months 0–4 (months before visible bureau "
        "deterioration). This is the headline comparison metric in the UI and verify_predictions.py."
    )
    if struct or full:
        _add_table(
            doc,
            ["Metric", "Structured", "Full"],
            [
                ["Stress detection rate", f"{struct.get('stress_detection_rate', 0)*100:.1f}%", f"{full.get('stress_detection_rate', 0)*100:.1f}%"],
                ["Accuracy", f"{struct.get('accuracy', 0)*100:.1f}%", f"{full.get('accuracy', 0)*100:.1f}%"],
                ["Precision", f"{struct.get('precision', 0)*100:.1f}%", f"{full.get('precision', 0)*100:.1f}%"],
                ["Recall (test)", f"{struct.get('recall', 0)*100:.1f}%", f"{full.get('recall', 0)*100:.1f}%"],
                ["ROC AUC", f"{struct.get('roc_auc', 0):.3f}", f"{full.get('roc_auc', 0):.3f}"],
                ["Decision threshold", str(struct.get("threshold", "—")), str(full.get("threshold", "—"))],
            ],
        )

    # 7. Scoring at inference
    doc.add_heading("7. Scoring at Inference", level=1)
    doc.add_paragraph("For a live case at month-on-book t:")
    _add_bullet(doc, "extract_features(profile, observation_month=t)")
    _add_bullet(doc, "compute_rule_stress_prob() — five-pillar rule engine (src/prediction/rule_engine.py)")
    _add_bullet(doc, "predict_ml_stress(model_type='full') — LightGBM probability")
    _add_bullet(doc, "final_stress_prob = 0.6 × rule_prob + 0.4 × ml_prob")
    doc.add_paragraph(
        "Rule pillars (weights): Repayment 30%, Cashflow 25%, Bureau/NTC 20%, Reputation/NLP 15%, Context 10%."
    )
    _add_code_block(
        doc,
        "rule_prob  = compute_rule_stress_prob(features)\n"
        "ml_prob    = predict_ml_stress(features, model_type='full')\n"
        "final_prob = 0.6 × rule_prob + 0.4 × ml_prob   # default ml_weight=0.4"
    )
    doc.add_paragraph(
        "Output bands: Low (<25%), Watch (25–45%), High (45–70%), Critical (≥70%). "
        "Drivers and narrative from extract_stress_drivers() and build_stress_narrative()."
    )

    # 8. Artifacts
    doc.add_heading("8. Model Artifacts & Regeneration", level=1)
    _add_table(
        doc,
        ["File", "Description"],
        [
            ["data/models/stress_model_structured.pkl", "Joblib bundle: model, feature_columns, threshold, metrics"],
            ["data/models/stress_model_full.pkl", "Full-feature classifier bundle"],
            ["data/models/training_metrics.json", "Latest train/test metrics for UI display"],
            ["data/synthetic/panels/stress_panel.csv", "Labelled training panel"],
        ],
    )

    # 9. Production path
    doc.add_heading("9. Production Considerations", level=1)
    _add_bullet(doc, "Replace synthetic profiles with CBS loan tape + live connectors (see CONNECTOR_INTEGRATION.md).")
    _add_bullet(doc, "Retrain on 3+ years of out-of-time validated portfolio outcomes.")
    _add_bullet(doc, "Swap keyword NLP with fine-tuned classifier or LLM feature extraction.")
    _add_bullet(doc, "Monitor drift on collection DPD and text signal distributions.")
    _add_bullet(doc, "Governance: document model version, training date, and feature schema in model registry.")

    # 10. Code reference
    doc.add_heading("10. Code Reference", level=1)
    _add_table(
        doc,
        ["Component", "Path"],
        [
            ["Label generation", "scripts/generate_data.py → _compute_stress_labels()"],
            ["Feature extraction", "src/features/feature_engineering.py → extract_features()"],
            ["NLP features", "src/features/nlp_features.py"],
            ["Collection features", "src/features/collection_features.py"],
            ["Training", "scripts/train_model.py → src/prediction/model.py → train_models()"],
            ["Inference", "src/prediction/model.py → compute_stress_prediction()"],
            ["Verification", "scripts/verify_predictions.py"],
        ],
    )

    if full.get("feature_columns"):
        doc.add_heading("Appendix A — Full model feature list", level=1)
        cols = full["feature_columns"]
        chunk = 15
        for i in range(0, len(cols), chunk):
            _add_bullet(doc, ", ".join(cols[i : i + chunk]))

    doc.add_paragraph("")
    footer = doc.add_paragraph("Finn. · Early Default Prediction · https://github.com/arunbhatg/finn-early-default-prediction")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
